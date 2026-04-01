from __future__ import annotations

import hashlib
import mimetypes
import os
import re
import subprocess
import tempfile
import unicodedata
from io import BytesIO
from pathlib import Path
from time import monotonic
from typing import Any, Protocol
from urllib import error as urlerror
from urllib import request as urlrequest
from urllib.parse import unquote, urlparse

from ..models.extraction import ExtractionResult, OCRText, PageExtraction
from ..utils.confidence import clamp_confidence

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SUPPORTED_MIME_TYPES = {"application/pdf", "image/jpeg", "image/png", DOCX_MIME_TYPE}


class ExtractionError(Exception):
    def __init__(
        self,
        message: str,
        *,
        code: str = "extraction_failed",
        status_code: int = 500,
        retriable: bool = False,
        details: dict[str, Any] | None = None,
    ) -> None:
        super().__init__(message)
        self.code = code
        self.status_code = status_code
        self.retriable = retriable
        self.details = details or {}


class PDFExtractor(Protocol):
    def extract_pages(self, payload: bytes) -> list[str]:
        ...


class DOCXExtractor(Protocol):
    def extract_text(self, payload: bytes) -> str:
        ...


class PDFPageRenderer(Protocol):
    def render_pages(self, payload: bytes) -> list[bytes]:
        ...


class OCRExtractor(Protocol):
    def extract(self, payload: bytes) -> OCRText:
        ...


class PypdfExtractor:
    def extract_pages(self, payload: bytes) -> list[str]:
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover
            raise ExtractionError(
                "PDF extraction dependency is missing",
                code="dependency_unavailable",
                status_code=500,
                retriable=False,
                details={"dependency": "pypdf"},
            ) from exc

        try:
            reader = PdfReader(BytesIO(payload))
            return [(page.extract_text() or "") for page in reader.pages]
        except Exception as exc:  # pragma: no cover
            raise ExtractionError(
                "could not parse PDF input",
                code="invalid_argument",
                status_code=400,
                retriable=False,
                details={"error": str(exc)},
            ) from exc


class TesseractOCRExtractor:
    def extract(self, payload: bytes) -> OCRText:
        try:
            import pytesseract
            from PIL import Image
        except ImportError as exc:  # pragma: no cover
            raise ExtractionError(
                "OCR dependencies are missing",
                code="dependency_unavailable",
                status_code=500,
                retriable=False,
                details={"dependencies": ["pytesseract", "Pillow"]},
            ) from exc

        try:
            image = Image.open(BytesIO(payload))
            text = pytesseract.image_to_string(image)
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        except Exception as exc:  # pragma: no cover
            raise ExtractionError(
                "OCR processing failed",
                code="dependency_unavailable",
                status_code=500,
                retriable=True,
                details={"error": str(exc)},
            ) from exc

        values: list[float] = []
        for raw in data.get("conf", []):
            try:
                score = float(raw)
            except (TypeError, ValueError):
                continue
            if score >= 0:
                values.append(score / 100.0 if score > 1 else score)

        confidence = sum(values) / len(values) if values else None
        return OCRText(
            text=text,
            confidence=confidence,
            diagnostics={"engine": "tesseract", "recognized_tokens": len(values)},
        )


class PopplerPDFRenderer:
    def render_pages(self, payload: bytes) -> list[bytes]:
        with tempfile.TemporaryDirectory(prefix="pdf-ocr-") as tmp_dir:
            pdf_path = Path(tmp_dir) / "input.pdf"
            output_prefix = Path(tmp_dir) / "page"
            pdf_path.write_bytes(payload)

            try:
                subprocess.run(
                    [
                        "pdftoppm",
                        "-png",
                        "-r",
                        "200",
                        str(pdf_path),
                        str(output_prefix),
                    ],
                    check=True,
                    capture_output=True,
                    text=True,
                )
            except FileNotFoundError as exc:  # pragma: no cover
                raise ExtractionError(
                    "PDF OCR renderer is unavailable",
                    code="dependency_unavailable",
                    status_code=500,
                    retriable=False,
                    details={"dependency": "pdftoppm"},
                ) from exc
            except subprocess.CalledProcessError as exc:  # pragma: no cover
                raise ExtractionError(
                    "PDF page rendering failed",
                    code="dependency_unavailable",
                    status_code=500,
                    retriable=True,
                    details={"stderr": (exc.stderr or "").strip()},
                ) from exc

            images = sorted(Path(tmp_dir).glob("page-*.png"))
            if not images:
                raise ExtractionError(
                    "PDF page rendering produced no images",
                    code="dependency_unavailable",
                    status_code=500,
                    retriable=True,
                )
            return [image_path.read_bytes() for image_path in images]


class PythonDocxExtractor:
    def extract_text(self, payload: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover
            raise ExtractionError(
                "DOCX extraction dependency is missing",
                code="dependency_unavailable",
                status_code=500,
                retriable=False,
                details={"dependency": "python-docx"},
            ) from exc

        try:
            document = Document(BytesIO(payload))
        except Exception as exc:  # pragma: no cover
            raise ExtractionError(
                "could not parse DOCX input",
                code="invalid_argument",
                status_code=400,
                retriable=False,
                details={"error": str(exc)},
            ) from exc

        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)


class ExtractionPipeline:
    def __init__(
        self,
        *,
        pdf_extractor: PDFExtractor | None = None,
        docx_extractor: DOCXExtractor | None = None,
        ocr_extractor: OCRExtractor | None = None,
        pdf_page_renderer: PDFPageRenderer | None = None,
    ) -> None:
        self._pdf_extractor = pdf_extractor or PypdfExtractor()
        self._docx_extractor = docx_extractor or PythonDocxExtractor()
        self._ocr_extractor = ocr_extractor or TesseractOCRExtractor()
        self._pdf_page_renderer = pdf_page_renderer or PopplerPDFRenderer()

    def extract_from_uri(self, storage_uri: str, mime_type: str | None = None) -> ExtractionResult:
        payload, file_path = _load_document_bytes(storage_uri)
        resolved_mime = _resolve_mime_type(mime_type, file_path)
        return self.extract_bytes(payload, resolved_mime)

    def extract_bytes(self, payload: bytes, mime_type: str) -> ExtractionResult:
        started = monotonic()
        mime_type = _resolve_mime_type(mime_type, None)

        raw_pages: list[str]
        ocr_used = False
        ocr_diag: dict[str, Any] = {}
        source = "pdf_text"
        confidences: list[float] = []

        if mime_type == "application/pdf":
            raw_pages = self._pdf_extractor.extract_pages(payload)
            if not raw_pages:
                raw_pages = [""]
            if _should_fallback_to_pdf_ocr(raw_pages):
                raw_pages, confidences, ocr_diag = self._extract_pdf_pages_with_ocr(payload)
                source = "ocr"
                ocr_used = True
        elif mime_type == DOCX_MIME_TYPE:
            source = "docx_text"
            raw_pages = [self._docx_extractor.extract_text(payload)]
        elif mime_type in {"image/jpeg", "image/png"}:
            source = "ocr"
            ocr_used = True
            ocr_text = self._ocr_extractor.extract(payload)
            raw_pages = [ocr_text.text]
            ocr_diag = dict(ocr_text.diagnostics)
            if ocr_text.confidence is not None:
                confidences = [clamp_confidence(ocr_text.confidence)]
        else:
            raise ExtractionError(
                f"unsupported mime_type: {mime_type}",
                code="invalid_argument",
                status_code=400,
                retriable=False,
                details={"supported_mime_types": sorted(SUPPORTED_MIME_TYPES)},
            )

        pages: list[PageExtraction] = []
        for idx, text in enumerate(raw_pages, start=1):
            normalized = _normalize_page_text(text)
            confidence = confidences[idx - 1] if len(confidences) >= idx else _heuristic_confidence(normalized)
            pages.append(
                PageExtraction(
                    page_number=idx,
                    text=normalized,
                    char_count=len(normalized),
                    confidence=round(confidence, 3),
                    source=source,
                )
            )

        full_text = "\n\f\n".join(page.text for page in pages)
        mean_confidence = 0.0 if not pages else round(sum(page.confidence for page in pages) / len(pages), 3)
        empty_pages = [page.page_number for page in pages if page.char_count == 0]
        diagnostics = {
            "ocr_used": ocr_used,
            "page_count": len(pages),
            "empty_pages": empty_pages,
            "char_count": sum(page.char_count for page in pages),
            "normalization": "unicode-nfkc + whitespace-collapse",
            "checksum_sha256": hashlib.sha256(payload).hexdigest(),
            "processing_ms": int((monotonic() - started) * 1000),
        }
        if ocr_diag:
            diagnostics["ocr"] = ocr_diag

        return ExtractionResult(
            mime_type=mime_type,
            text=full_text,
            pages=pages,
            confidence=mean_confidence,
            diagnostics=diagnostics,
        )

    def _extract_pdf_pages_with_ocr(self, payload: bytes) -> tuple[list[str], list[float], dict[str, Any]]:
        rendered_pages = self._pdf_page_renderer.render_pages(payload)
        texts: list[str] = []
        confidences: list[float] = []
        page_diagnostics: list[dict[str, Any]] = []

        for image_payload in rendered_pages:
            ocr_text = self._ocr_extractor.extract(image_payload)
            texts.append(ocr_text.text)
            confidences.append(
                clamp_confidence(ocr_text.confidence) if ocr_text.confidence is not None else 0.0
            )
            page_diagnostics.append(dict(ocr_text.diagnostics))

        return (
            texts,
            confidences,
            {
                "engine": "tesseract",
                "pdf_fallback": True,
                "rendered_page_count": len(rendered_pages),
                "page_diagnostics": page_diagnostics,
            },
        )


def _resolve_mime_type(mime_type: str | None, file_path: str | None) -> str:
    value = (mime_type or "").strip().lower()
    if value == "image/jpg":
        value = "image/jpeg"
    if not value and file_path:
        guessed, _ = mimetypes.guess_type(file_path)
        value = (guessed or "").lower()
    if value not in SUPPORTED_MIME_TYPES:
        raise ExtractionError(
            f"unsupported mime_type: {value or 'unknown'}",
            code="invalid_argument",
            status_code=400,
            retriable=False,
            details={"supported_mime_types": sorted(SUPPORTED_MIME_TYPES)},
        )
    return value


def _load_document_bytes(storage_uri: str) -> tuple[bytes, str]:
    parsed = urlparse(storage_uri)
    if parsed.scheme in {"http", "https"}:
        request = urlrequest.Request(storage_uri, method="GET")
        try:
            with urlrequest.urlopen(request, timeout=20) as response:
                return response.read(), parsed.path or storage_uri
        except urlerror.HTTPError as exc:
            raise ExtractionError(
                "storage_uri http request failed",
                code="invalid_argument",
                status_code=400,
                retriable=False,
                details={"storage_uri": storage_uri, "status_code": exc.code},
            ) from exc
        except urlerror.URLError as exc:
            raise ExtractionError(
                "storage_uri http request failed",
                code="dependency_unavailable",
                status_code=502,
                retriable=True,
                details={"storage_uri": storage_uri, "error": str(exc.reason)},
            ) from exc

    if parsed.scheme not in {"", "file"}:
        raise ExtractionError(
            f"unsupported storage_uri scheme: {parsed.scheme}",
            code="invalid_argument",
            status_code=400,
            retriable=False,
            details={"storage_uri": storage_uri},
        )

    path = unquote(parsed.path) if parsed.scheme == "file" else storage_uri
    if not path:
        raise ExtractionError(
            "storage_uri path is empty",
            code="invalid_argument",
            status_code=400,
            retriable=False,
        )
    if not os.path.exists(path):
        raise ExtractionError(
            "storage_uri file does not exist",
            code="invalid_argument",
            status_code=400,
            retriable=False,
            details={"path": path},
        )

    with open(path, "rb") as file:
        return file.read(), path


def _normalize_page_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKC", text)
    normalized = normalized.replace("\u00a0", " ")
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\f", "\n")

    compact_lines: list[str] = []
    previous_blank = False
    for line in normalized.split("\n"):
        line = re.sub(r"[ \t]+", " ", line).strip()
        if not line:
            if not previous_blank:
                compact_lines.append("")
            previous_blank = True
            continue
        previous_blank = False
        compact_lines.append(line)

    while compact_lines and compact_lines[-1] == "":
        compact_lines.pop()
    while compact_lines and compact_lines[0] == "":
        compact_lines.pop(0)
    return "\n".join(compact_lines)


def _should_fallback_to_pdf_ocr(raw_pages: list[str]) -> bool:
    normalized_pages = [_normalize_page_text(text) for text in raw_pages]
    return all(not page for page in normalized_pages)


def _heuristic_confidence(text: str) -> float:
    chars = len(text)
    if chars == 0:
        return 0.05
    if chars < 40:
        return 0.55
    if chars < 160:
        return 0.72
    return 0.88


def _clamp_confidence(value: float) -> float:
    return clamp_confidence(value)
